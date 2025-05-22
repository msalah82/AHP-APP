
import streamlit as st
import numpy as np
import pandas as pd
import json
from docx import Document
import matplotlib.pyplot as plt
from io import BytesIO
from tempfile import NamedTemporaryFile

st.set_page_config(page_title="AHP Tool", layout="centered")
st.title("📊 AHP Decision Support Tool")

# Define top-level input fields first to avoid duplicate key issues

# Load metadata for criteria/alternatives if present in session
session_meta_criteria = st.session_state.get("session_metadata_criteria", "Cost, Quality, Time")
session_meta_alternatives = st.session_state.get("session_metadata_alternatives", "Option A, Option B, Option C")

# Show text inputs using metadata as default
criteria_input = st.text_input("Enter criteria (comma-separated)", session_meta_criteria, key="criteria_input")
alternative_input = st.text_input("Enter alternatives (comma-separated)", session_meta_alternatives, key="alternative_input")

criteria = [c.strip() for c in criteria_input.split(",") if c.strip()]
alternatives = [a.strip() for a in alternative_input.split(",") if a.strip()]

# Handle session upload and apply state without rerun
st.sidebar.title("💾 Save/Load Session")
if st.sidebar.button("🔄 Save Session"):
    session_data = {k: v for k, v in st.session_state.items() if k not in ["criteria_input", "alternative_input"]}
        session_data["session_metadata_criteria"] = criteria_input
    session_data["session_metadata_alternatives"] = alternative_input

    with open("ahp_session.json", "w") as f:
        json.dump(session_data, f)
    with open("ahp_session.json", "rb") as f:
        st.sidebar.download_button("📥 Download Session", f, file_name="ahp_session.json")


uploaded_file = st.sidebar.file_uploader("📤 Load Session", type="json")
if uploaded_file and "session_applied" not in st.session_state:
    session_data = json.load(uploaded_file)
    for k, v in session_data.items():
        try:
            if k in ["criteria_input", "alternative_input"]:
            continue
        elif k == "session_metadata":
            continue

        if k == "session_metadata_criteria":
            st.session_state["session_metadata_criteria"] = v
            continue
        if k == "session_metadata_alternatives":
            st.session_state["session_metadata_alternatives"] = v
            continue
        st.session_state[k] = v
        except Exception:
            st.warning(f"⚠️ Skipped invalid session item: {k}")
    st.session_state.session_applied = True

def get_pairwise_matrix(items, title, session_key):
    st.subheader(title)
    size = len(items)
    matrix = np.ones((size, size))
    if session_key not in st.session_state:
        st.session_state[session_key] = {}
    for i in range(size):
        for j in range(i + 1, size):
            label = f"{items[i]} vs {items[j]}"
            key = f"{session_key}_{label}"
            val = st.slider(label, 1/9.0, 9.0, st.session_state[session_key].get(label, 1.00), step=0.01, format="%.2f", key=key)
            st.session_state[session_key][label] = val
            matrix[i, j] = val
            matrix[j, i] = 1 / val
    return matrix

def calculate_weights(matrix):
    eigvals, eigvecs = np.linalg.eig(matrix)
    max_index = np.argmax(np.real(eigvals))
    weights = np.real(eigvecs[:, max_index])
    return weights / np.sum(weights)

def calculate_consistency_ratio(matrix):
    n = len(matrix)
    eigvals, _ = np.linalg.eig(matrix)
    max_eigval = np.max(np.real(eigvals))
    CI = (max_eigval - n) / (n - 1) if n > 1 else 0
    RI_dict = {1: 0.0, 2: 0.0, 3: 0.58, 4: 0.90, 5: 1.12, 6: 1.24, 7: 1.32, 8: 1.41, 9: 1.45, 10: 1.49}
    RI = RI_dict.get(n, 1.49)
    return (CI / RI) if RI != 0 else 0

if len(criteria) >= 2 and len(alternatives) >= 2:
    criteria_matrix = get_pairwise_matrix(criteria, "Criteria Comparison", "criteria")
    criteria_CR = calculate_consistency_ratio(criteria_matrix)
    st.write(f"**Criteria Consistency Ratio:** {criteria_CR:.2f}")
    st.dataframe(pd.DataFrame(criteria_matrix, index=criteria, columns=criteria))

    criteria_weights = calculate_weights(criteria_matrix)
    alt_scores = np.zeros(len(alternatives))
    alt_weights_dict = {}
    final_matrix_log = {}

    for idx, crit in enumerate(criteria):
        key = f"alt_{crit}"
        alt_matrix = get_pairwise_matrix(alternatives, f"{crit} Comparison", key)
        CR = calculate_consistency_ratio(alt_matrix)
        st.write(f"**CR for {crit}:** {CR:.2f}")
        st.dataframe(pd.DataFrame(alt_matrix, index=alternatives, columns=alternatives))
        weights = calculate_weights(alt_matrix)
        alt_weights_dict[crit] = weights
        final_matrix_log[crit] = alt_matrix.tolist()
        alt_scores += weights * criteria_weights[idx]

    results_df = pd.DataFrame({
        "Alternative": alternatives,
        "Final Score": alt_scores
    }).sort_values("Final Score", ascending=False).reset_index(drop=True)

    st.subheader("🏆 Final Ranking")
    winner = results_df.iloc[0]
    st.success(f"🥇 **Winner: {winner['Alternative']}** with score {winner['Final Score']:.4f}")

    st.subheader("📊 Bar Chart of Scores")
    fig, ax = plt.subplots()
    ax.bar(results_df["Alternative"], results_df["Final Score"])
    st.pyplot(fig)

    if st.button("⬇️ Export to Excel"):
        output = BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            pd.DataFrame(criteria_matrix, index=criteria, columns=criteria).to_excel(writer, sheet_name="Criteria Matrix")
            for crit in criteria:
                pd.DataFrame(final_matrix_log[crit], index=alternatives, columns=alternatives).to_excel(writer, sheet_name=f"{crit}")
            results_df.to_excel(writer, sheet_name="Final Ranking", index=False)
        st.download_button("📥 Download Excel", output.getvalue(), file_name="AHP_Results.xlsx")

    if st.button("📝 Export to Word"):
        doc = Document()
        doc.add_heading("📊 AHP Decision Support Report", 0)
        doc.add_heading("Criteria", level=1)
        for i, c in enumerate(criteria):
            doc.add_paragraph(f"{i+1}. {c} (Weight: {criteria_weights[i]:.4f})")
        doc.add_heading("Alternatives", level=1)
        for a in alternatives:
            doc.add_paragraph(f"- {a}")
        doc.add_heading("Final Scores", level=1)
        for i, row in results_df.iterrows():
            alt = row["Alternative"]
            score = row["Final Score"]
            text = f"🥇 {alt}: {score:.4f}" if i == 0 else f"{alt}: {score:.4f}"
            doc.add_paragraph(text)
        doc.add_paragraph("Generated using AHP Web Tool with Streamlit.")
        with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp.seek(0)
            st.download_button("📄 Download Word Report", tmp.read(), file_name="AHP_Report.docx")
else:
    st.info("Enter at least 2 criteria and 2 alternatives to start.")
